from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict, Any
import os
from pathlib import Path
from agent_project.config import WORKSPACE_ROOT

def _resolve_path(path_str: str) -> Path:
    """Resolves path relative to WORKSPACE_ROOT."""
    path = Path(path_str)
    if path.is_absolute():
        return path.resolve()
    else:
        return (Path(WORKSPACE_ROOT) / path).resolve()

def _add_sections_to_doc(doc, sections: List[Dict[str, Any]]):
    """Helper to add sections to a document object."""
    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        
        if section_type == "heading":
            level = section.get("level", 1)
            heading = doc.add_heading(content, level=level)
            alignment = section.get("alignment", "left")
            if alignment == "center":
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "right":
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        elif section_type == "paragraph":
            para = doc.add_paragraph()
            run = para.add_run(content)
            
            # Apply formatting
            if section.get("bold"):
                run.bold = True
            if section.get("italic"):
                run.italic = True
            
            # Alignment
            alignment = section.get("alignment", "left")
            if alignment == "center":
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "right":
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == "justify":
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        elif section_type == "bullet_list":
            items = content if isinstance(content, list) else [content]
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
        
        elif section_type == "numbered_list":
            items = content if isinstance(content, list) else [content]
            for item in items:
                doc.add_paragraph(item, style='List Number')
        
        elif section_type == "table":
            # content should be List[List[str]]
            if isinstance(content, list) and len(content) > 0:
                rows = len(content)
                cols = len(content[0]) if content[0] else 1
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for i, row_data in enumerate(content):
                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        row.cells[j].text = str(cell_data)

def create_word_document(
    file_path: str,
    title: str = None,
    sections: List[Dict[str, Any]] = None,
    author: str = "AI Agent",
    append: bool = False,
    writing_plan: str = None
) -> Dict[str, str]:
    """
    Cria ou EDITA ITERATIVAMENTE um documento .docx profissional.
    
    **PROTOCOLO DE PLANEJAMENTO (DOCUMENTOS LONGOS - 10+ páginas):**
    1. **Primeira Chamada (Inicialização):**
       - `append=False`, `writing_plan="1. Intro\\n2. Desenvolvimento\\n3. Conclusão"`
       - Cria o documento E salva o plano mestre em `.filename.docx.plan`
    
    2. **Chamadas Subsequentes (Iterações):**
       - `append=True`, `writing_plan="COMPLETANDO: Seção 2"`
       - Adiciona novas seções ao documento existente
       - O sistema mostra: plano original + preview do conteúdo atual
    
    **REQUISITO:** Estruture o conteúdo em uma lista de objetos `sections` (títulos, parágrafos, bullets, tabelas).
    **BENEFÍCIO:** Evita limites de token, mantém coerência e previne "amnésia" em documentos longos.
    """
    from datetime import datetime
    
    target_path = _resolve_path(file_path)
    plan_path = target_path.parent / f".{target_path.name}.plan"
    
    print(f"\n--- {'Appending to' if append else 'Creating'} Word Document: {target_path} ---")
    
    try:
        # === SISTEMA DE PLANEJAMENTO ===
        if append and target_path.exists():
            # Modo APPEND: Mostrar contexto + plano
            print(f"\n[Planning System] APPEND MODE")
            
            # Ler documento existente para contexto
            existing_doc = Document(target_path)
            num_paragraphs = len(existing_doc.paragraphs)
            num_sections = len(existing_doc.sections)
            
            # Pegar últimos 3 parágrafos como contexto
            last_paragraphs = existing_doc.paragraphs[-3:] if num_paragraphs >= 3 else existing_doc.paragraphs
            context_preview = "\n".join([p.text[:200] for p in last_paragraphs if p.text.strip()])
            
            print(f"[Context] Documento tem {num_paragraphs} parágrafos, {num_sections} seções")
            print(f"\n[Context] Últimos parágrafos:\n{context_preview}\n")
            
            # Mostrar plano mestre (se existe)
            if plan_path.exists():
                with open(plan_path, 'r', encoding='utf-8') as f:
                    plan_content = f.read()
                print(f"[Master Plan]\n{plan_content}\n")
            
            # Atualizar plano com progresso
            if writing_plan and plan_path.exists():
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                with open(plan_path, 'a', encoding='utf-8') as f:
                    f.write(f"[{timestamp}] {writing_plan}\n")
            
            # Adicionar seções ao documento existente
            doc = existing_doc
            if sections:
                _add_sections_to_doc(doc, sections)
        
        else:
            # Modo CREATE: Novo documento
            doc = Document()
            
            # Set document properties
            if author:
                doc.core_properties.author = author
            if title:
                doc.core_properties.title = title
            
            # Add title if provided
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add sections
            if sections:
                _add_sections_to_doc(doc, sections)
            
            # Salvar plano mestre (se fornecido)
            if writing_plan:
                plan_path.parent.mkdir(parents=True, exist_ok=True)
                with open(plan_path, 'w', encoding='utf-8') as f:
                    f.write("MASTER PLAN:\n")
                    f.write(f"{writing_plan}\n\n")
                    f.write("PROGRESS LOG:\n")
                print(f"[Planning System] Master plan saved to: {plan_path}")
        
        # AVISO: Documento grande sem planejamento
        if not append and not writing_plan and sections and len(sections) > 20:
            print(f"\n⚠️  [LARGE DOCUMENT DETECTED] ({len(sections)} sections)")
            print(f"   RECOMENDAÇÃO: Use o sistema de planejamento!")
            print(f"   Exemplo:")
            print(f"     1. create_word_document('{target_path.name}', sections=parte1, writing_plan='1. Intro\\n2. Body\\n3. Conclusão')")
            print(f"     2. create_word_document('{target_path.name}', sections=parte2, append=True, writing_plan='COMPLETANDO: Body')")
        
        # Ensure directory exists
        target_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(target_path)
        
        action = "appended to" if append else "created"
        return {
            "status": "Success",
            "message": f"Word document {action} successfully: {target_path}",
            "file_path": str(target_path),
            "plan_file": str(plan_path) if writing_plan else None
        }
        
    except Exception as e:
        return {
            "status": "Error",
            "message": f"Failed to create/append Word document: {str(e)}"
        }

def edit_word_document(
    file_path: str,
    operations: List[Dict[str, Any]]
) -> Dict[str, str]:
    """
    Edits an existing Word document.
    
    Args:
        file_path: Path to the .docx file
        operations: List of operations to perform
        
    Operation types:
    1. Append content:
       { "action": "append", "section": { ... same as create_word_document section ... } }
       
    2. Replace text:
       { "action": "replace", "target": "old text", "replacement": "new text" }
       
    Returns:
        Dict with status
    """
    target_path = _resolve_path(file_path)
    print(f"\n--- Editing Word Document: {target_path} ---")
    
    try:
        if not target_path.exists():
            return {"status": "Error", "message": f"File not found: {target_path}"}
            
        doc = Document(target_path)
        
        for op in operations:
            action = op.get("action")
            
            if action == "append":
                section = op.get("section")
                if section:
                    _add_sections_to_doc(doc, [section])
            
            elif action == "replace":
                target = op.get("target")
                replacement = op.get("replacement")
                if target and replacement:
                    # Simple replacement in paragraphs
                    for para in doc.paragraphs:
                        if target in para.text:
                            para.text = para.text.replace(target, replacement)
                    # Simple replacement in tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if target in para.text:
                                        para.text = para.text.replace(target, replacement)
        
        doc.save(target_path)
        
        return {
            "status": "Success",
            "message": f"Word document edited successfully: {target_path}",
            "file_path": str(target_path)
        }
        
    except Exception as e:
        return {
            "status": "Error",
            "message": f"Failed to edit Word document: {str(e)}"
        }
